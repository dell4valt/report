import os
import random
import sys
import unittest
from pathlib import Path
from docx.shared import Cm

# Получаем абсолютный путь к основной папке проекта
# и добавляем его в переменную среды PYTHONPATH
project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))

from report.core import Report


class TestReport(unittest.TestCase):
    def test_create_report(self):
        report = Report()
        report.add_paragraph("Hello, world!")
        filename = f"temp_{random.randrange(1, 10**3):05}.docx"
        report.save(filename)
        self.assertTrue(True)  # Просто проверяем, что код не падает
        os.remove(filename)

    def test_add_paragraph(self):
        report = Report()
        report.add_paragraph("Hello, world!")
        self.assertEqual(len(report.doc.paragraphs), 1)

    def test_add_heading(self):
        report = Report()
        report.add_heading("Hello, world!", level=1)
        self.assertEqual(len(report.doc.paragraphs), 1)
        self.assertEqual(
            report.doc.paragraphs[0].style.name, report.STYLES["heading"][1]
        )

    def test_save(self):
        report = Report()
        filename = f"temp_{random.randrange(1, 10**3):05}.docx"
        report.save(filename)
        self.assertTrue(os.path.exists(filename))
        os.remove(filename)

    def test_set_last_paragraph_style(self):
        report = Report()
        report.add_paragraph("Hello, world!")
        report.set_last_paragraph_style("Heading 1")
        self.assertEqual(report.doc.paragraphs[-1].style.name, "Heading 1")


if __name__ == "__main__":
    unittest.main()
