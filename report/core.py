"""Библиотека предоставляет класс и набор полезных функция для
формирования технического отчета в формате Microsoft Word
с помощью библиотеки python-docx. Таких как:

    * Report - класс для создания отчетов в файле Microsoft Word
    * get_xls_sheet_quantity - функция возвращающая количество листов в файле Excel
    * set_table_columns_width - установка ширины колонок в таблице docx.Table
    * set_table_style - установка стиля текста в ячейках всей таблицы docx.Table
    * set_table_rows_style - установка стиля текста в указанных строках таблицы docx.Table
"""

import os
import random
from pathlib import Path
from importlib import resources

import pandas as pd
from docx import Document
from docx.table import Table
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt
from report.utils import insert_row_numbers_in_df


class Report:
    """Класс для создания отчетов в файле Microsoft Word.
    Позволяет создавать отчеты с текстовой информацией, таблицами и графиками
    и вставлять их в существующий документ или создавать новый.

    Args:
        file_path (str, optional): путь к файлу, в который добавляется отчет.
    Если не указан или файл не существует, отчет создается на основе шаблона по умолчанию
    DEFAULT_TEMPLATE. По умолчанию None.
    """

    # Путь к файлу шаблона отчета по умолчанию
    DEFAULT_TEMPLATE = str(resources.files("report").joinpath("templates/template.docx"))

    STYLES = {
        "table": {
            "title": "Т-название",
            "text": "Т-таблица",
            "text_heading": "Т-таблица-заголовок",
            "footer": "Т-примечание",
            "table": "Table Grid",
        },
        "figure": {
            "title": "Р-название",
            "fig": "Р-рисунок",
        },
        "heading": {
            1: "Heading 1",
            2: "Heading 2",
            3: "Heading 3",
            4: "Heading 4",
            5: "Heading 5",
            6: "Heading 6",
        },
    }

    TEXT = {
        "parameter": "Параметр",
        "figure": "Рисунок",
        "table": "Таблица",
    }

    def __init__(self, file_path=None) -> None:
        """Инициализирует новый экземпляр класса Report.

        Args:
            file_path (str, optional): Путь к файлу, в который добавляется отчет.
        Если не указан или файл не существует, отчет создается на основе шаблона по умолчанию.
        По умолчанию None.
        """
        self.doc = Document

        if file_path:
            print(f"\nРезультаты расчетов будут добавлены к файлу {file_path}")

            # Пытаемся загрузить существующий файл
            # и вставляем разрыв страницы
            self.set_template(file_path)
        else:
            # Создаем новый документ на основе стандартного шаблона
            self._init_empty_report()

    def _init_empty_report(self) -> None:
        """Метод инициализирует пустой отчет на основе стандартного шаблона
        и удаляет первый параграф.
        """
        self.set_template(self.DEFAULT_TEMPLATE)
        self.remove_paragraph()

    def add_paragraph(self, text: str, style=None) -> None:
        """Добавляет в документ абзац с указанным текстом и стилем.

        Args:
          text (str): Текст который необходимо добавить в документ в виде абзаца
          style (str): Стиль применяемый к добавляемому абзацу.
        """
        self.doc.add_paragraph(text, style=style)

    def add_heading(self, text, level=1) -> None:
        """Добавляет в документ абзац заголовка заданного уровня (1-6).

        Args:
          text (str): Текст заголовка
          level (int): Уровень заголовка (возможны значения 1-6). По умолчанию 1.
        """

        if not 1 <= level <= 6:
            print(
                "Уровень заголовка не может быть меньше 1 и больше 6, "
                f"заголовок '{text}' будет записан обычным текстом."
            )
            self.doc.add_paragraph(text)
            return

        self.doc.add_paragraph(text, style=self.STYLES["heading"][level])

    def merge_table_cells(
        self,
        table,
        start_row: int,
        end_row: int,
        start_column: int,
        end_column: int,
        value: str = "",
        style: str = "Т-таблица",
    ) -> None:
        """Метод объединяет ячейки таблицы и при необходимости
        вставляет текст в объединенные ячейки.

        Args:
            table (Table): Таблица.
            start_row (int): Номер строки начала объединения.
            end_row (int): Номер строки конца объединения.
            start_column (int): Номер столбца начала объединения.
            end_column (int): Номер столбца конца объединения.
            value (str, optional): Текст, который необходимо вставить в объединенную ячейку.
            style (str, optional): Стиль текста в объединенной ячейке.
        """
        col_num = len(table.columns)
        row_num = len(table.rows)

        # Проверка значений
        if start_row >= row_num or end_row >= row_num:
            raise ValueError("Значение start_row или end_row больше количества строк в таблице.")
        if start_column >= col_num or end_column >= col_num:
            raise ValueError("Значение start_column или end_column больше количества столбцов в таблице.")

        if (start_row > end_row or start_column > end_column) or (
            start_row < 0 or end_row < 0 or start_column < 0 or end_column < 0
        ):
            raise ValueError(
                f"Неверные значения для объединения ячейки: "
                f"start_row={start_row}, end_row={end_row}, "
                f"start_column={start_column}, end_column={end_column}"
            )

        # Объединяем ячейки
        table.cell(start_row, start_column).merge(table.cell(end_row, end_column))

        # Устанавливаем значения ячейке и стиль текста
        if value:
            table.cell(start_row, start_column).text = value
            for paragraph in table.cell(start_row, start_column).paragraphs:
                paragraph.style = style

    def get_table_cell_value(self, table, row: int, column: int) -> str:
        """Возвращает текст из ячейки таблицы.

        Args:
            table (Table): Таблица.
            row (int): Номер строки.
            column (int): Номер столбца.
        """
        if not table:
            raise ValueError("Таблица не должна быть пустой.")

        if row < 0 or column < 0:
            raise ValueError("Значение row или column не может быть отрицательным.")

        if row >= len(table.rows) or column >= len(table.columns):
            raise IndexError("Значение row или column выходит за границы таблицы.")

        cell = table.cell(row, column)
        if not cell:
            raise ValueError("Cell is null")

        return cell.text

    @staticmethod
    def set_table_cell_value(table, row: int, column: int, value: str, style: str = "Т-таблица") -> None:
        """Устан

        Args:
            table (Table): Таблица.
            row (int): Номер строки.
            column (int): Номер столбца.
            value (str): Текст, который необходимо вставить в ячейку.
            style (str, optional): Стиль текста в ячейке. По умолчанию "Т-таблица".
        """
        if not table:
            raise ValueError("Таблица не должна быть пустой.")

        if row < 0 or column < 0:
            raise ValueError("Значение row или column не может быть отрицательным.")

        if row >= len(table.rows) or column >= len(table.columns):
            raise IndexError("Значение row или column выходит за границы таблицы.")

        table.cell(row, column).text = value

        if style:
            set_table_cell_style(table, row, column, style=style)

    def insert_df_to_table(
        self,
        df,
        title=None,
        footer_text=None,
        col_names=None,
        col_widths=None,
        col_format=None,
        table_style=STYLES["table"]["table"],
        text_style=STYLES["table"]["text"],
        first_row_table_style=STYLES["table"]["text_heading"],
        row_names=None,
        rows_idx=False,
    ) -> None:
        """Метод вставляет Pandas DataFrame в файл отчета как отформатированную
        таблицу с возможностью указания заголовков для колонок, стилей таблицы и
        текста, ширины столбцов.

        Args:
            df (pandas.DataFrame): DataFrame содержащий данные, которые необходимо вставить
                в таблицу документа Word.
            title (str, optional): Заголовок (название) таблицы, которая будет вставлена в документ.
                По умолчанию None.
            footer_text (str, optional): Текст примечания, который будет добавлен под таблицей.
                По умолчанию None.
            col_names (tuple, list, optional): Переопределяет заголовки колонок в таблице.
                По умолчанию заголовки соответствуют названию колонок в DataFrame.
            col_widths (tuple, optional): Переопределяет ширину колонок таблицы по порядку.
                По умолчанию None.
            col_format (tuple, optional): Стиль форматирования значений для каждого столбца таблицы.
                Форматы соответствуют f-строке (пример: ":g", ":g", ":.2f").
                По умолчанию None.
            table_style (str, optional): Название устанавливаемого стиля таблицы.
                Стиль должен присутствовать в файле шаблона отчета.
                По умолчанию STYLES["table"]["table"].
            text_style (str, optional): Название устанавливаемого основного стиля текста в таблице.
                Стиль должен присутствовать в файле шаблона отчета.
                По умолчанию STYLES["table"]["text"].
            first_row_table_style (str, optional): Название устанавливаемого стиля строки
                заголовков таблицы. Стиль должен присутствовать в файле шаблона отчета.
                По умолчанию STYLES["table"]["text_heading"].
            row_names (list, tuple, str, optional): Названия строк для добавления в первый столбец.
                По умолчанию None.
            rows_idx (bool, optional): Флаг для добавления нумерации строк. По умолчанию False.

        Returns:
            docx.table.Table: Экземпляр таблицы с произведенными изменениями.
        """
        doc = self.doc

        # Проверка на тип данных и пустой DataFrame
        if not isinstance(df, pd.DataFrame):
            print("\nОшибка! Для вставки таблиц необходимо передать Pandas.DataFrame.")
            print(f"Таблица '{title}' не будет вставлена в отчет.\n")
            return

        if df.empty:
            print("\nОшибка! Для вставки таблицы необходимо передать не пустой Pandas.DataFrame.")
            print(f"Таблица '{title}' не будет вставлена в отчет.\n")
            return

        # Вставляем заголовок таблицы
        if title:
            doc.add_paragraph(f"{self.TEXT['table']} — {title}", style=self.STYLES["table"]["title"])

        if rows_idx:
            df = insert_row_numbers_in_df(df, "№")

        # Количество строк и столбцов в таблице
        rows = df.shape[0]
        columns = df.shape[1]

        # Учитываем дополнительную колонку для row_names
        add_row_names = isinstance(row_names, (list, tuple, str))
        if add_row_names:
            columns += 1

        # Добавляем таблицу в документ
        table = doc.add_table(rows + 1, columns, style=table_style)

        # Получаем доступ к ячейкам экземпляра таблицы для
        # увеличения производительности, все последующие
        # операции производим с ячейками а не с экземпляром таблицы
        cells = table._cells

        # Устанавливаем 1ю строку заголовков
        for column_idx, column_name in enumerate(df.columns):
            # Сдвиг на 1, если если присутствуют названия строк
            header_idx = column_idx + (1 if add_row_names else 0)
            if col_names:
                try:
                    cells[header_idx].text = str(col_names[column_idx])
                except IndexError:
                    raise IndexError(
                        f"Ошибка в задании количества заголовков к таблице, задано {len(col_names)} заголовков, а в таблице {columns} колонок."
                    )

            else:
                cells[header_idx].text = str(column_name)

        # Добавляем заголовок для первой колонки, если row_names указано
        if add_row_names:
            cells[0].text = self.TEXT["parameter"]
            for row_idx, name in enumerate(row_names):
                cells[(row_idx + 1) * columns].text = str(name)

        # Записываем данные df в таблицу
        for row_idx in range(rows):
            for column_idx in range(df.shape[1]):
                cell_value = df.iat[row_idx, column_idx]

                cell_idx = (row_idx + 1) * columns + column_idx + (1 if add_row_names else 0)

                # Проверка на количество форматов
                if col_format and len(col_format) != df.shape[1]:
                    raise ValueError(
                        "Длина списка форматов не соответствует количеству столбцов в таблице. "
                        f"Заданное количество форматов: {len(col_format)}, количество столбцов в таблице: {df.shape[1]}"
                    )

                # Если задан список стилей для форматирования текста
                # устанавливаем формат для значения каждой ячейки
                # иначе просто записываем строку значения в ячейку
                if col_format and isinstance(cell_value, (float, int)):
                    s = f"{{{col_format[column_idx]}}}"
                    cells[cell_idx].text = s.format(cell_value)
                else:
                    cells[cell_idx].text = str(cell_value)

        if col_widths:
            self._set_table_columns_width(table, col_widths)

        self._set_table_style(table, text_style, first_row_table_style)

        # Записываем служебный параграф после таблиц
        if footer_text:
            doc.add_paragraph(footer_text, style=self.STYLES["table"]["footer"])
        else:
            doc.add_paragraph("", style=self.STYLES["table"]["footer"])

        return table

    @staticmethod
    def insert_table_row(table, index: int):
        """
        Вставляет новую строку в таблицу по указанному индексу.

        Args:
            table: Таблица docx, в которую нужно вставить строку
            index: Индекс, по которому будет вставлена строка (0 = начало таблицы)

        Returns:
            Вставленная строка
        """
        # Проверяем допустимость индекса
        row_count = len(table.rows)
        if index < 0 or index > row_count:
            raise ValueError(f"Индекс {index} вне диапазона [0, {row_count}]")

        # Добавляем новую строку в конец
        new_row = table.add_row()

        # Если нужно вставить в конец, работа уже выполнена
        if index == row_count:
            return new_row

        # Иначе перемещаем строки снизу вверх, начиная с последней
        for i in range(row_count, index, -1):
            # Исходная строка (которую нужно переместить вниз)
            src_row = table.rows[i - 1]
            # Целевая строка (куда нужно переместить)
            dst_row = table.rows[i]

            # Копируем содержимое и форматирование каждой ячейки
            for j in range(len(table.columns)):
                if j < len(src_row.cells) and j < len(dst_row.cells):
                    # Копирование текста
                    dst_cell = dst_row.cells[j]
                    src_cell = src_row.cells[j]

                    # Безопасное копирование текста
                    if src_cell.paragraphs:
                        # Очищаем все параграфы, кроме первого
                        for p in list(dst_cell.paragraphs)[1:]:
                            p._element.getparent().remove(p._element)

                        # Копируем текст первого параграфа
                        if dst_cell.paragraphs and src_cell.paragraphs:
                            # Безопасное копирование текста
                            dst_cell.paragraphs[0].text = src_cell.paragraphs[0].text or ""

                            # Копирование стиля (проверка на None)
                            if hasattr(src_cell.paragraphs[0], "style") and src_cell.paragraphs[0].style:
                                try:
                                    dst_cell.paragraphs[0].style = src_cell.paragraphs[0].style
                                except:
                                    pass  # Игнорируем ошибки со стилями

                            # Копируем дополнительные параграфы
                            for src_para in src_cell.paragraphs[1:]:
                                dst_para = dst_cell.add_paragraph()
                                dst_para.text = src_para.text or ""

                                # Копирование стиля (с проверкой)
                                if hasattr(src_para, "style") and src_para.style:
                                    try:
                                        dst_para.style = src_para.style
                                    except:
                                        pass  # Игнорируем ошибки со стилями
                    else:
                        dst_cell.text = ""

                    # Безопасное копирование свойств ячейки
                    try:
                        v_align = src_cell._tc.get("vAlign")
                        if v_align is not None:  # Проверка на None
                            dst_cell._tc.set("vAlign", v_align)
                    except:
                        pass  # Игнорируем ошибку, если свойство не может быть скопировано

                    # Безопасное копирование tcPr
                    try:
                        if src_cell._tc.tcPr is not None:
                            dst_cell._tc.tcPr = deepcopy(src_cell._tc.tcPr)
                    except:
                        pass  # Игнорируем ошибку, если свойство не может быть скопировано

        # Очищаем строку по указанному индексу
        target_row = table.rows[index]

        # Получаем шаблонную строку для форматирования
        template_row_index = min(index + 1, row_count - 1) if index < row_count - 1 else max(0, index - 1)
        template_row = table.rows[template_row_index]

        # Очищаем содержимое, сохраняя базовое форматирование
        for j, cell in enumerate(target_row.cells):
            if j < len(template_row.cells):
                # Очищаем содержимое параграфов, но оставляем хотя бы один
                for p in list(cell.paragraphs)[1:]:
                    try:
                        p._element.getparent().remove(p._element)
                    except:
                        pass

                if cell.paragraphs:
                    cell.paragraphs[0].text = ""

                # Копируем базовые свойства ячейки (безопасно)
                try:
                    if template_row.cells[j]._tc.tcPr is not None:
                        cell._tc.tcPr = deepcopy(template_row.cells[j]._tc.tcPr)
                except:
                    pass  # Игнорируем ошибки копирования свойств

        return target_row

    def insert_table_column(self, table: Table, index: int, values: list[str] = [], width: float = 0.4) -> None:
        """Добавляет столбец в таблицу.

        Args:
            table: Таблица python-docx
            index: Индекс позиции, на которую нужно добавить новый столбец
            values: Список значений, которые нужно вставить в новый столбец
            width: Ширина столбца в сантиметрах. По умолчанию 0.4
        """
        table.add_column(Cm(width))
        last_col_idx = len(table.rows[0].cells) - 1

        self.move_table_column(table, last_col_idx, index)

        if values:
            for i, val in enumerate(values):
                self.set_table_cell_value(table, i, index, val)

    @staticmethod
    def move_table_column(table, from_index, to_index):
        """
        Перемещает колонку в таблице table с позиции from_index на позицию to_index.
        Индексация с 0.
        """
        if from_index < 0 or to_index < 0:
            raise ValueError("Значение from_index или to_index не может быть отрицательным.")

        if from_index >= len(table.rows[0].cells) or to_index >= len(table.rows[0].cells):
            raise IndexError("Значение from_index или to_index выходит за границы таблицы.")

        for row in table.rows:
            cells = list(row._tr)  # XML-элементы ячеек <w:tc>
            cell = cells.pop(from_index)
            cells.insert(to_index, cell)

            # Перезаписываем структуру строки
            row._tr.clear()
            for c in cells:
                row._tr.append(c)

    @staticmethod
    def insert_table_second_row_header(
        table,
        merge_vertical_cols: list[int] = [0, 1],
        merge_horizontal_start_col: int = 2,
        merged_horizontal_text: str = "",
    ):
        """
        Добавляет строку в начало таблицы и объединяет ячейки по правилам.

        Args:
            table: Таблица python-docx
            merge_vertical_cols (list): Индексы столбцов, которые нужно объединить по вертикали (строки 1 и 2). По умолчанию [0, 1]
            merge_horizontal_start_col (int): Индекс столбца, с которого нужно начать горизонтальное объединение в первой строке. По умолчанию 2
            merged_horizontal_text (str): Текст для объединённой горизонтальной ячейки. По умолчанию ""
        """

        # Вспомогательная функция для установки текста ячейки
        def set_cell_text(cell, text: str):
            # Удаляем все параграфы
            for p in cell.paragraphs:
                p.clear()
            # Добавляем один параграф с нужным текстом
            cell.text = text.strip()

        # Добавляем новую строку в конец и копируем строки вверх
        table.add_row()
        for i in range(len(table.rows) - 1, 0, -1):
            for j in range(len(table.columns)):
                table.cell(i, j).text = table.cell(i - 1, j).text

        # Объединение вертикальных ячеек (строки 0 и 1)
        for col in merge_vertical_cols:
            if col < len(table.columns):
                cell_top = table.cell(0, col)
                cell_bottom = table.cell(1, col)

                # Сохраняем текст верхней ячейки
                top_text = cell_top.text.strip()

                # Очистка нижней ячейки (останется мусор)
                cell_bottom.text = ""

                # Объединениям и устанавливаем текст
                merged_cell = cell_top.merge(cell_bottom)
                set_cell_text(merged_cell, top_text)

        # Объединение горизонтальных ячеек в строке 0
        if merge_horizontal_start_col < len(table.columns) - 1:
            row = table.rows[0]
            start_cell = row.cells[merge_horizontal_start_col]
            end_cell = row.cells[-1]

            # Очистим текст в объединяемых ячейках, кроме первой
            for i in range(merge_horizontal_start_col + 1, len(table.columns)):
                row.cells[i].text = ""

            merged_cell = start_cell.merge(end_cell)
            set_cell_text(merged_cell, merged_horizontal_text)

    def insert_mpl_figure(self, chart, title="", dpi=200, width=16.5) -> None:
        """Метод вставляет график Matplotlib.plt в документ, предварительно
        сохранив его во временный файл. Устанавливает стили, добавляет заголовок
        и затем удаляет временный файл.

        Args:
        chart (Matplotlib.plt): график который необходимо вставить в документ
        title (str): заголовок графика
        dpi (int): разрешение с которым будет вставлен график. По умолчанию 200
        width (float): ширина графика в сантиметрах
        """
        doc = self.doc

        # Временная директория в которую сохраняется график
        temp_dir = "TEMP"

        # Название временного файла графика
        filename = f"temp_{random.randrange(1, 10**3):03}.png"

        # Путь сохранения графика в файл
        file_path = f"{temp_dir}/{filename}"
        # Создаем временную папку, если она отсутствует
        Path(file_path).parents[0].mkdir(parents=True, exist_ok=True)
        # Сохраняем график в файл
        chart.savefig(file_path, dpi=dpi)
        # Вставляем график в документ отчёта
        self.insert_picture(file_path, title=title, width=width)

        # Пытаемся удалить файл графика
        try:
            os.remove(file_path)
        except FileNotFoundError:
            pass

    def insert_picture(self, file_path, title=None, width=16.5) -> None:
        """Метод вставляет изображение в документ Word.

        Args:
            file_path (str): Путь к картинке.
            width (float): Ширина картинки в сантиметрах.
        """
        doc = self.doc
        doc.add_picture(file_path, width=Cm(width))
        # Устанавливаем стиль графика
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.style = self.STYLES["figure"]["fig"]
        if title:
            doc.add_paragraph(f"{self.TEXT['figure']} — {title}", style=self.STYLES["figure"]["title"])

    def insert_page_break(self) -> None:
        """Метод вставляет в документ Word разрыв страницы
        в том месте где вызывается.
        """

        doc = self.doc
        paragraphs = doc.paragraphs
        run = paragraphs[-1].add_run()
        run.add_break(WD_BREAK.PAGE)

    def _set_table_columns_width(self, table, col_widths: tuple) -> None:
        """Метод устанавливает ширину столбцов в таблице docx
        поочередно проходя по каждой ячейке таблицы.

        Args:
            table (docx.table.Table): Таблица в документе docx.
            col_widths (tuple): Кортеж, содержащий ширину столбцов
        в сантиметрах по порядку.
        """
        set_table_columns_width(table, col_widths)

    def _set_table_style(
        self,
        table,
        style=STYLES["table"]["text"],
        first_row_style=STYLES["table"]["text_heading"],
    ) -> None:
        """Метод проходил по всем ячейкам таблицы table и устанавливает
        заданный стиль style параграфов в таблице. При желании можно
        указать стиль для заголовков таблицы (1-ая строка).

        Args:
            table (docx.table.Table): Таблица в документе docx
            style (str): Название устанавливаемого стиля. Стиль должен присутствовать
            в файле шаблона отчета. По умолчанию "Т-таблица".
            first_row_style (str): Названия стиля для первой строки таблицы.
            (строка заголовков). По умолчанию None.
        """
        set_table_style(table, style, first_row_style)

    def _set_table_rows_style(self, table, rows=(0, 1), style=STYLES["table"]["text_heading"]) -> None:
        """Метод предназначен для установки заданного стиля для указанных строк таблицы.

        Args:
            table (docx.table.Table): Таблица в документе docx
            style (str): Название устанавливаемого стиля. Стиль должен присутствовать
            в файле шаблона отчета. По умолчанию "Т-таблица-заголовок".
        """
        set_table_rows_style(table, rows, style)

    def set_last_paragraph_style(self, style: str) -> None:
        """Метод устанавливает стиль последнего по порядку
        параграфа в документе Word.

        Args:
            style (str): Название стиля в документе
        (стиль обязательно должен присутствовать в документе отчёта).
        """
        doc = self.doc
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.style = style

    def set_template(self, template: str) -> None:
        """Метод позволяет установить шаблон документа Microsoft Word.
        Если указанный файл шаблон не найден, используется шаблон по умолчанию
        ("report/assets/template.docx").

        Args:
            template (str): Путь к файлу шаблону
        """

        # Проверка типа аргумента template
        if not isinstance(template, str):
            print("Ошибка! Аргумент template должен быть строкой.")
            self._init_empty_report()
            return

        template_path = Path(template)

        # Проверка наличия файла и его расширения
        if template_path.is_file():
            if template_path.suffix.lower() in [".docx", ".doc"]:
                self.doc = Document(template)
            else:
                print(f"Ошибка! Файл шаблона {template} должен иметь расширение .docx или .doc.")
                print("Отчет будет сохранен с файлом шаблона по умолчанию.")
                self._init_empty_report()
        else:
            print(f"Файл {template_path} не обнаружен.")
            print("Отчет будет сохранен с файлом шаблона по умолчанию.")
            if os.path.exists(self.DEFAULT_TEMPLATE):
                self._init_empty_report()
            else:
                raise FileExistsError(f"Файл шаблона по умолчанию: {self.DEFAULT_TEMPLATE} не обнаружен.")

    def remove_paragraph(self, index=-1) -> None:
        """Метод удаляет параграф из файла документа. По умолчанию удаляется текущий параграф.

        Args:
            index (int): Индекс удаляемого параграфа, -1 — текущий параграф, -2 предыдущий и т.д.
        По умолчанию: -1 (текущий)
        """
        # -1 будет текущий параграф, -2 будет предыдущий
        try:
            paragraph = self.doc.paragraphs[index]

            # Удаляем параграф из документа
            p = paragraph._element
            p.getparent().remove(p)
        except IndexError:
            print(f"Ошибка, не удалось удалить параграф с индексом {index}. Вероятно этот параграф не существует.")

    def save(self, filename: str) -> None:
        """Метод сохраняет отчёт в файл Microsoft Word по заданному пути.

        Args:
            filename (str): Путь к файлу в который сохранить отчёт
        """
        report_path = Path(filename)

        # Создаем папку до пути файла, если не существует
        report_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            self.doc.save(report_path)
            print("\n" + "-" * 74)
            print(f"Файл: {report_path} успешно сохранён!")
            print("-" * 74 + "\n")
        except PermissionError:
            print("\nОшибка! Не удалось сохранить файл. Проверьте возможность записи файла по указанному пути.")
            print("Возможно записываемый файл уже существует и открыт.")

            self._handle_save_error(filename)

    def _handle_save_error(self, path: str) -> None:
        temp_path = str(Path(path).parent) + f"/temp_{random.randrange(1, 10**3):03}.docx"
        inp = input(f"\nСохранить во временный файл {temp_path} (y/n)? ")
        if inp.lower() == "y":
            self.save(temp_path)
        elif inp.lower() == "n":
            print("\nОтчет не сохранен.")
            return
        else:
            self._handle_save_error(path)


def set_table_columns_width(table, col_widths: tuple) -> None:
    """Функция устанавливает ширину столбцов в таблице docx
    поочередно проходя по каждой ячейке таблицы.

    Args:
        table (docx.table.Table): Таблица в документе docx.
        col_widths (tuple): Кортеж, содержащий ширину столбцов
    в сантиметрах по порядку.
    """
    # Получаем доступ к ячейкам таблицы из соображений производительности
    # и считаем количество колонок и строк
    cells = table._cells
    columns = len(table.columns)
    rows = len(table.rows)

    if columns != len(col_widths):
        print("\nВнимание количество заданных столбцов не совпадает с количеством столбцов в таблице.")
        print(f"В таблице — {columns}, задано — {len(col_widths)}")

    for row_idx in range(rows):
        for column_idx in range(columns):
            # Номер ячейки по порядку
            cell_n = column_idx + row_idx * columns

            # Устанавливаем ширину столбцов.
            # Если ячеек заданных ширин меньше чем столбцов
            # устанавливаем последнее успешное значение
            try:
                success_width = Cm(col_widths[column_idx])
                cells[cell_n].width = success_width

            except IndexError:
                cells[cell_n].width = success_width


def set_table_style(table, style="Т-таблица", first_row_style="Т-таблица-заголовок") -> None:
    """Функция проходит по всем ячейкам таблицы table и устанавливает
    заданный стиль style параграфов в таблице. При желании можно
    указать стиль для заголовков таблицы (1-ая строка).

    Args:
        table (docx.table.Table): Таблица в документе docx
        style (str): Название устанавливаемого стиля. Стиль должен присутствовать
        в файле шаблона отчета. По умолчанию "Т-таблица".
        first_row_style (str): Названия стиля для первой строки таблицы.
        (строка заголовков). По умолчанию None.
    """
    # Получаем доступ к ячейкам таблицы из соображений производительности
    # и считаем количество колонок и строк
    cells = table._cells
    columns = len(table.columns)
    rows = len(table.rows)

    for row_idx in range(rows):
        for column_idx in range(columns):
            # Номер ячейки по порядку
            cell_n = column_idx + row_idx * columns

            # Если задан стиль первой строки устанавливаем
            # иначе проходим по всем ячейкам и устанавливаем
            # основной стиль таблицы
            for paragraph in cells[cell_n].paragraphs:
                paragraph.style = style

            if row_idx == 0 and first_row_style:
                for paragraph in cells[cell_n].paragraphs:
                    try:
                        paragraph.style = first_row_style
                    except TypeError:
                        print(
                            "Ошибка установки стиля первой строки, "
                            f"заголовок: {paragraph.text}, стиль: {first_row_style}"
                        )


def set_table_font_size(table, font_size=10) -> None:
    """Метод проходит по всем ячейкам таблицы table и устанавливает
    заданный размер шрифта font_size параграфов в таблице.

    Args:
        table (docx.table.Table): Таблица в документе docx
        font_size (int): Размер шрифта. По умолчанию 10.
    """
    # Получаем доступ к ячейкам таблицы из соображений производительности
    # и считаем количество колонок и строк
    cells = table._cells
    columns = len(table.columns)
    rows = len(table.rows)

    for row_idx in range(rows):
        for column_idx in range(columns):
            # Номер ячейки по порядку
            cell_n = column_idx + row_idx * columns

            # Устанавливаем размер шрифта
            for paragraph in cells[cell_n].paragraphs:
                paragraph.style.font.size = Pt(font_size)


def set_table_rows_style(table, rows=(0, 1), style="Т-таблица-заголовок") -> None:
    """Метод предназначен для установки заданного стиля для указанных строк таблицы.

    Args:
        table (docx.table.Table): Таблица в документе docx
        style (str): Название устанавливаемого стиля. Стиль должен присутствовать
        в файле шаблона отчета. По умолчанию "Т-таблица-заголовок".
    """
    cells = table._cells
    columns = len(table.columns)

    for row_idx in rows:
        for column_idx in range(columns):
            # Номер ячейки по порядку
            cell_n = column_idx + row_idx * columns

            # Если задан стиль первой строки устанавливаем
            # иначе проходим по всем ячейкам и устанавливаем
            # основной стиль таблицы
            for paragraph in cells[cell_n].paragraphs:
                paragraph.style = style


def set_table_cell_style(table, row, column, style="Т-таблица") -> None:
    """Функция устанавливает заданный стиль style параграфов в ячейке таблицы table.
    Args:
        table (docx.table.Table): Таблица в документе docx
        row (int): Номер строки
        column (int): Номер колонки
        style (str): Название устанавливаемого стиля. Стиль должен присутствовать
        в файле шаблона отчета. По умолчанию "Т-таблица".
    """

    cell = table.cell(row, column)
    for paragraph in cell.paragraphs:
        paragraph.style = style
